"""
Text extraction from uploaded reports: PDF (native + scanned/OCR), DOCX, images,
and plain text. Every backend is optional and imported lazily so the core
pipeline still runs if a library is missing -- the function reports what it used.
"""
import os


def _from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _from_docx(path):
    import docx  # python-docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _from_pdf(path):
    """Try native text via pdfplumber/PyMuPDF; OCR pages that have ~no text."""
    text_parts = []
    used_ocr = False
    # Prefer PyMuPDF (fitz) if present, else pdfplumber.
    pages_text = []
    try:
        import fitz
        doc = fitz.open(path)
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
    except Exception:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    pages_text.append(page.extract_text() or "")
        except Exception:
            pages_text = []

    for i, pt in enumerate(pages_text):
        if pt and len(pt.strip()) > 20:
            text_parts.append(pt)
        else:
            ocr = _ocr_pdf_page(path, i)
            if ocr:
                used_ocr = True
                text_parts.append(ocr)
    if not pages_text:  # nothing extracted at all -> OCR whole doc
        ocr_all = _ocr_pdf_all(path)
        if ocr_all:
            used_ocr = True
            text_parts.append(ocr_all)
    return "\n".join(text_parts), used_ocr


def _ocr_pdf_page(path, page_index):
    try:
        import fitz, pytesseract
        from PIL import Image
        import io
        doc = fitz.open(path)
        page = doc[page_index]
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        doc.close()
        return pytesseract.image_to_string(img)
    except Exception:
        return ""


def _ocr_pdf_all(path):
    out = []
    i = 0
    while True:
        t = _ocr_pdf_page(path, i)
        if not t and i > 0:
            break
        if t:
            out.append(t)
        i += 1
        if i > 50:
            break
    return "\n".join(out)


def _from_image(path):
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(path))


def extract_text(path):
    """Return (text, meta) where meta documents the method/quality.
    Raises RuntimeError with a helpful message if extraction is impossible."""
    ext = os.path.splitext(path)[1].lower()
    meta = {"ext": ext, "method": None, "ocr": False, "ok": True, "warning": ""}
    try:
        if ext in (".txt", ".md"):
            text = _from_txt(path); meta["method"] = "plaintext"
        elif ext in (".docx",):
            text = _from_docx(path); meta["method"] = "docx"
        elif ext in (".doc",):
            raise RuntimeError("Legacy .doc not supported; please save as .docx or PDF.")
        elif ext == ".pdf":
            text, used_ocr = _from_pdf(path)
            meta["method"] = "pdf+ocr" if used_ocr else "pdf-native"
            meta["ocr"] = used_ocr
        elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"):
            text = _from_image(path); meta["method"] = "image-ocr"; meta["ocr"] = True
        else:
            raise RuntimeError(f"Unsupported file type: {ext}")
    except ImportError as e:
        meta["ok"] = False
        meta["warning"] = (f"Missing library for {ext} ({e}). "
                           f"Install requirements.txt to enable this format.")
        return "", meta
    text = (text or "").strip()
    if len(text) < 80:
        meta["warning"] = ("Very little text extracted; the file may be a "
                           "low-quality scan. OCR results can be unreliable.")
    return text, meta
